from docx import Document

def create_dpa():
    doc = Document()
    doc.add_heading('Data Processing Agreement', 0)

    doc.add_heading('Preamble', 1)
    doc.add_paragraph('This Data Processing Agreement (DPA) supplements the Master Services Agreement between the Data Controller (Client) and the Data Processor (Provider).')

    doc.add_heading('Article 1: Scope of Processing', 1)
    doc.add_paragraph('The Processor shall process personal data solely on behalf of the Controller and strictly in accordance with documented instructions. The data processed includes names, emails, and IP addresses.')

    doc.add_heading('Article 2: Data Security', 1)
    doc.add_paragraph('The Processor must implement appropriate technical and organizational measures to protect personal data against accidental deletion or unauthorized access, as outlined in Article 1.')

    doc.add_heading('Article 3: Data Breach Notification', 1)
    doc.add_paragraph('In the event of a personal data breach, the Processor shall notify the Controller without undue delay, and in no event later than 72 hours after having become aware of it.')

    doc.save('test_dpa.docx')
    print("Sukces! Plik test_dpa.docx czeka w folderze głównym.")

if __name__ == "__main__":
    create_dpa()