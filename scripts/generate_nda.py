from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_nda():
    doc = Document()
    
    # Tytuł
    title = doc.add_heading('Mutual Non-Disclosure Agreement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This Agreement is entered into by and between Party A and Party B.')
    
    # Artykuł 1
    doc.add_heading('Article 1: Definitions', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('1.1 "Confidential Information" ').bold = True
    p1.add_run('means any data or information, oral or written, treated as confidential.')
    
    # Artykuł 2 (z referencją do Artykułu 1)
    doc.add_heading('Article 2: Obligations', level=1)
    doc.add_paragraph('2.1 The Receiving Party shall not disclose the Confidential Information defined in Article 1.')
    
    # Zagnieżdżona sekcja
    doc.add_heading('Section 2.1.1: Exceptions', level=2)
    doc.add_paragraph('Information that is public knowledge is not considered confidential.')
    
    doc.save('test_nda.docx')
    print("Test NDA generated successfully as 'test_nda.docx'!")

if __name__ == "__main__":
    create_test_nda()