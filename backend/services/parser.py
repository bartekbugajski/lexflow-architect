from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

from models.legal_objects import Clause


@dataclass(frozen=True)
class _HeaderMatch:
    header_type: str
    number: str
    remainder: str
    level: int


class LegalDocParser:
    """
    Parses a .docx into a hierarchy of Clause objects.

    Design goals:
    - Deterministic splitting into atomic clauses using common legal headers.
    - Preserve nested structure via parent_id.
    - Capture paragraph-level formatting metadata needed for reconstruction.
    """

    _ARTICLE_RE = re.compile(
        r"^\s*Article\s+(?P<num>(?:\d+|[IVXLC]+))\b(?P<rest>.*)$",
        re.IGNORECASE,
    )
    _SECTION_RE = re.compile(
        r"^\s*Section\s+(?P<num>\d+(?:\.\d+)*)\b(?P<rest>.*)$",
        re.IGNORECASE,
    )

    # Finds in-body references like "Article 3" or "Section 2.1"
    _REF_RE = re.compile(
        r"\b(?:(Article)\s+(\d+|[IVXLC]+)|(Section)\s+(\d+(?:\.\d+)*))\b",
        re.IGNORECASE,
    )

    def parse_path(self, docx_path: str) -> List[Clause]:
        doc = Document(docx_path)
        return self._parse_document(doc)

    def parse_bytes(self, docx_bytes: bytes) -> List[Clause]:
        doc = Document(io.BytesIO(docx_bytes))
        return self._parse_document(doc)

    def _parse_document(self, doc: Document) -> List[Clause]:
        clauses: List[Clause] = []
        # Header stack only (Article/Section/etc). Preamble is kept separate so it
        # never becomes an unintended parent for later headers.
        stack: List[Tuple[int, Clause]] = []
        preamble_clause: Optional[Clause] = None

        for para in doc.paragraphs:
            raw_text = para.text or ""
            text = raw_text.strip()
            if not text:
                continue

            header = self._match_header(text)
            formatting = self._extract_paragraph_formatting(para)

            if header:
                clause = self._new_clause(
                    title=self._normalize_header_title(text),
                    parent_id=self._parent_id_for_level(stack, header.level),
                    header=header,
                    header_formatting=formatting,
                )
                clauses.append(clause)

                # Maintain nesting stack
                while stack and stack[-1][0] >= header.level:
                    stack.pop()
                stack.append((header.level, clause))
                continue

            if stack:
                current = stack[-1][1]
            else:
                if preamble_clause is None:
                    preamble_clause = self._new_clause(
                        title="Preamble",
                        parent_id=None,
                        header=_HeaderMatch(
                            header_type="preamble",
                            number="",
                            remainder="",
                            level=1,
                        ),
                        header_formatting={},
                    )
                    clauses.append(preamble_clause)
                current = preamble_clause
            self._append_paragraph(current, text=text, formatting=formatting)

        # Assign a stable, sequential order_index to each clause based on
        # the order they were discovered in the source document.
        for idx, clause in enumerate(clauses):
            clause.order_index = idx

        return clauses

    def _new_clause(
        self,
        *,
        title: Optional[str],
        parent_id: Optional[str],
        header: _HeaderMatch,
        header_formatting: Dict[str, Any],
    ) -> Clause:
        return Clause(
            id=str(uuid.uuid4()),
            title=title,
            text="",
            parent_id=parent_id,
            metadata={
                "header": {
                    "type": header.header_type,
                    "number": header.number,
                    "remainder": header.remainder.strip(),
                    "level": header.level,
                },
                "header_formatting": header_formatting,
                "paragraphs": [],
            },
        )

    def _append_paragraph(self, clause: Clause, *, text: str, formatting: Dict[str, Any]) -> None:
        paragraphs = clause.metadata.setdefault("paragraphs", [])
        if isinstance(paragraphs, list):
            paragraphs.append({"text": text, "formatting": formatting})
        else:
            clause.metadata["paragraphs"] = [{"text": text, "formatting": formatting}]

        existing = clause.text.strip()
        clause.text = f"{existing}\n\n{text}".strip() if existing else text

    def _match_header(self, text: str) -> Optional[_HeaderMatch]:
        m = self._ARTICLE_RE.match(text)
        if m:
            num = m.group("num")
            rest = (m.group("rest") or "").strip(" \t-–—:.")
            return _HeaderMatch(
                header_type="article",
                number=num,
                remainder=rest,
                level=1,
            )

        m = self._SECTION_RE.match(text)
        if m:
            num = m.group("num")
            rest = (m.group("rest") or "").strip(" \t-–—:.")
            depth = len(num.split("."))
            level = 2 + max(depth - 1, 0)
            return _HeaderMatch(
                header_type="section",
                number=num,
                remainder=rest,
                level=level,
            )

        # Heuristic fallback: treat Word heading styles as headers.
        # This is intentionally conservative: only Heading 1-4.
        # (We still preserve formatting metadata for reconstruction.)
        return None

    def _normalize_header_title(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _parent_id_for_level(self, stack: List[Tuple[int, Clause]], level: int) -> Optional[str]:
        if not stack:
            return None
        for existing_level, clause in reversed(stack):
            if existing_level < level:
                return clause.id
        return None

    def _extract_paragraph_formatting(self, para) -> Dict[str, Any]:
        alignment = None
        try:
            if para.alignment is not None:
                alignment = WD_ALIGN_PARAGRAPH(para.alignment).name
        except Exception:
            alignment = str(para.alignment) if para.alignment is not None else None

        bold = False
        italic = False
        for run in getattr(para, "runs", []) or []:
            if run.bold is True:
                bold = True
            if run.italic is True:
                italic = True
            if bold and italic:
                break

        style_name = None
        try:
            style_name = para.style.name if para.style is not None else None
        except Exception:
            style_name = None

        return {
            "bold": bold,
            "italic": italic,
            "alignment": alignment,
            "style": style_name,
        }

    def extract_references(self, clauses: List[Clause]) -> List[Dict[str, str]]:
        """
        Optional helper: extract intra-document references based on "Article X"/"Section Y.Z"
        mentions inside clause bodies. Returned list is stable and suitable for Neo4j UNWIND.
        """
        key_to_id: Dict[str, str] = {}
        for c in clauses:
            if not c.title:
                continue
            key = self._header_key(c.title)
            if key:
                key_to_id[key] = c.id

        refs: List[Dict[str, str]] = []
        for c in clauses:
            for ref_key in self._find_reference_keys(c.text or ""):
                target_id = key_to_id.get(ref_key)
                if not target_id or target_id == c.id:
                    continue
                refs.append({"src": c.id, "dst": target_id, "kind": ref_key})
        return refs

    def _header_key(self, title: str) -> Optional[str]:
        m = self._ARTICLE_RE.match(title)
        if m:
            return f"article {m.group('num').lower()}"
        m = self._SECTION_RE.match(title)
        if m:
            return f"section {m.group('num').lower()}"
        return None

    def _find_reference_keys(self, text: str) -> List[str]:
        keys: List[str] = []
        for m in self._REF_RE.finditer(text):
            if m.group(1) and m.group(2):
                keys.append(f"article {m.group(2).lower()}")
            elif m.group(3) and m.group(4):
                keys.append(f"section {m.group(4).lower()}")
        return keys

