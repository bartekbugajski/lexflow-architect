from __future__ import annotations

import io
from typing import Any, List

from docx import Document as DocxDocument  # type: ignore[import-untyped]

from services.graph_service import GraphService


class DocumentNotFoundError(Exception):
    """Raised when a requested Document node does not exist."""


class DocumentExporter:
    @staticmethod
    def export_to_docx(document_id: str, graph_service: GraphService) -> io.BytesIO:
        """
        Export a Document and its root Clause nodes to a .docx buffer.

        The Cypher query matches the Document and all directly contained Clauses:

            MATCH (d:Document {id: $doc_id})-[:CONTAINS]->(c:Clause)
            RETURN d, c
        """
        records: List[Any] = graph_service.run_read(
            """
            MATCH (d:Document {id: $doc_id})-[:CONTAINS]->(c:Clause)
            RETURN d AS document, c AS clause
            ORDER BY c.id
            """,
            doc_id=document_id,
        )

        if not records:
            raise DocumentNotFoundError(f"Document not found for id: {document_id}")

        docx = DocxDocument()

        for record in records:
            clause_node = record["clause"]
            clause_props = dict(clause_node)

            title = clause_props.get("title") or ""
            text = clause_props.get("text") or ""

            if title:
                docx.add_heading(title, level=2)
            if text:
                docx.add_paragraph(text)

        buffer = io.BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        return buffer

